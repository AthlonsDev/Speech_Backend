api_key = 'AIzaSyCoOw-xfas4_coDHFbo3nKyPHtSDk7HdGU'
import time
from google import genai
from docx import Document
from aws_client import upload_doc
import json
import os
import io


models = [
    "gemini-2.5-flash",
    "gemma-3-27b-it", 
    "gemini-2.0-flash",
    "gemini-3-flash-preview"
]
client = genai.Client(api_key=api_key)
client.models.list()

def get_date():
    return time.strftime("%d-%m-%Y")

def convert_to_doc(text, output_filename, type):
    doc = Document() #document object
    doc.add_heading('Transcription', level=0) #Title of the document
    doc.add_heading('document info', level=2) #Subtitle of the document
    doc.add_paragraph(f'date: {get_date()} \ntype of document: {type}') #Intro paragraph
    doc.add_heading('Transcribed Text', level=2) #Heading for transcribed text
    doc.add_paragraph(text) #First paragraph of the document
    doc.add_heading('Summary', level=1) #Summary heading
    # doc.add_paragraph(doc_assistant(text)) #Summary paragraph
    target_stream = io.BytesIO()
    doc.save(target_stream) #save doc to a stream
    target_stream.seek(0) #reset stream position to the beginning
    doc_bytes = target_stream.getvalue() #get the bytes of the document from the stream

    upload_doc(io.BytesIO(doc_bytes), output_filename) #upload doc to s3 using the bytes stream

    return output_filename

text=" The mute muffled the high tones of the horn. The gold ring fits only a pierced ear. The old pan was covered with hard fudge. Watch the log float in the wide river. The node on the stock of wheat grew daily. The heap of fallen leaves was set on fire. Right fast, if you want to finish early. His shirt was clean, but one button was gone. The barrel of beer was a brew of malt and hops. Tin cans are absent from store shelves."
# convert_to_doc(text, "Doc-Test", "transcription")


# output_filename = f"Transcription - {get_date()}.docx"
def doc_assistant(text):
    prompt = f"Summarize the following text in concise paragraph:\n\n{text}"
    response = client.models.generate_content(
        model=models[1],
        contents=prompt,
        config={
            'response_mime_type': 'text/plain',
        }
    )
    print("Assistant Response:", response.candidates[0].content.parts[0].model_dump()['text'])
    # return response.candidates[0].content.parts[0].model_dump()['text']
    summary = response.candidates[0].content.parts[0].model_dump()['text']

    return summary

meeting_struct = {
    'date': str,
    'attendees': list[str],
    'objectives': list[str],
    'action_items': list[str],
    'timeline': list[str],
    'summary': list[str],
    'conclusion': list[str]
}


def meetings_assistant(user_prompt: str , title):
    print("Generating meeting minutes...")
    prompt = f"You are an AI assistant that helps users with creating meetings minutes. Generate a text with information on timing and information based on transcription: \
    {user_prompt}. follow this structure {meeting_struct}, in objective and timeline try to highlight who said what. Keep the responses brief and to the point. Respond in JSON format only."
    response = client.models.generate_content(
        model=models[1],
        contents=prompt,
        config={
            'response_mime_type': 'text/plain',
        }
    )
    # print("Assistant Response:", response.candidates[0].content.parts[0].model_dump()['text'])
    text = response.candidates[0].content.parts[0].model_dump()['text']

    convert_text_to_json(text, title)

def meeting_doc_assistant(data, title):
    print("Generating meeting minutes document...")
    doc = Document() #document object
    doc.add_heading('Transcription', level=0) #Title of the document
    doc.add_heading('Meeting Info', level=0) #Subtitle of the document
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid' 
    table.cell(0, 0).text = f"date: {data['date']}"
    table.cell(0, 1).text = get_date()
    table.cell(1, 0).text = 'Time'
    table.cell(1, 1).text = '10:00 AM'
    table.cell(2, 0).text = 'Attendees'
    table.cell(2, 1).text = ", ".join(data['attendees'])
    doc.add_heading('Meeting Objectives', level=0) #Heading for transcribed text

    rows = len(data['objectives'])
    print(rows)
    table = doc.add_table(rows=rows, cols=1)
    table.style = 'Table Grid'
    for i in range(rows):
        print(data['objectives'][i])
        table.cell(i, 0).text = f'{data["objectives"][i]}'

    doc.add_heading('Meeting Timeline', level=0) #Summary heading

    rows = len(data['timeline'])
    table = doc.add_table(rows=rows, cols=1)
    table.style = 'Table Grid'
    for i in range(rows):
        print(data['timeline'][i])
        table.cell(i, 0).text = f'{data["timeline"][i]}'

    doc.add_heading('Meeting Summary', level=0) #Summary heading
    doc.add_paragraph(data['summary'][0]) #Summary paragraph

    output_filename = f"Meeting_{title}.docx"
    doc_path = doc.save(output_filename) #save doc locally and store path in variable
    upload_doc(doc_path, output_filename) #upload doc to s3
    os.remove(output_filename) #remove local copy of doc once uploaded to s3


    return output_filename


def convert_text_to_json(text, title):
    # convert the text to json format
    print("Converting text to JSON...")
    if text.startswith("```json") and text.endswith("```"):
        json_text = text[7:-3]  # Remove the ```json and ``` tags
        try:
            data = json.loads(json_text)
            data = {k: v for k, v in data.items() if k in meeting_struct}
            print("Extracted JSON data:", data)
            print(type(data))
            meeting_doc_assistant(data, title)
        except json.JSONDecodeError as e:
            print(f"Error decoding JSON: {e}")
            return None
    else: print("Text does not contain valid JSON format.")
    return None

# meetings_assistant(text, "Meeting_1")


# meeting_doc_assistant(data)
# meetings_assistant(text)